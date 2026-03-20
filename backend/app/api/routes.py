from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.schemas.query_schema import QueryRequest
from app.services.retrieval_service import search_documents
from app.services.rag_service import generate_answer
from fastapi import Depends
from app.models.document import Document
from fastapi import UploadFile, File
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.document import Document
import pdfplumber
from docx import Document as DocxDocument
import io

router = APIRouter()

@router.post("/ask")
def ask_question(request: QueryRequest, db: Session = Depends(get_db)):
    docs = search_documents(db, request.question)
    answer = generate_answer(request.question, docs)

    return {
        "question": request.question,
        "answer": answer,
        "documents_used": len(docs)
    }

def extract_text(file: UploadFile, content: bytes):
    filename = file.filename.lower()

    if filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    elif filename.endswith(".pdf"):
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    elif filename.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])

    else:
        return content.decode("utf-8", errors="ignore")


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):

    existing_doc = db.query(Document).filter(
        Document.title == file.filename
    ).first()

    if existing_doc:
        return {"message": "Document already exists"}
    content = await file.read()

    text = extract_text(file, content)

    # Remove problematic characters
    text = text.replace("\x00", "")

    if not text.strip():
        return {"error": "No readable text found in document"}

    doc = Document(
        title=file.filename,
        content=text
    )

    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {"message": "Document uploaded successfully"}

# @router.post("/upload")
# async def upload_document(
#     file: UploadFile = File(...),
#     db: Session = Depends(get_db)
# ):
#     content = await file.read()
#     text = content.decode("utf-8", errors="ignore")

#     doc = Document(
#         title=file.filename,
#         content=text
#     )

#     db.add(doc)
#     db.commit()
#     db.refresh(doc)

#     return {"message": "Document uploaded successfully"}